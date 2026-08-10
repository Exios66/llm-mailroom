import functools
import structlog
import time
from pathlib import Path
from typing import Any

from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver

from graph.state import DocumentState
from graph.routing import (
    after_classify,
    after_retry_classify,
    after_extraction,
    after_retry_extraction,
    after_boss,
    after_human_review,
)
from observability.tracing import pipeline_trace, traced_node, observation
from schemas.manifest import DocumentManifest, PipelineStage
from pipeline.config import get_confidence_thresholds
from pipeline.bins import (
    inbox_dir,
    processing_dir,
    classified_dir,
    review_dir,
    failed_dir,
    archive_dir,
    manifests_dir,
    ensure_dirs,
    claim_file,
    move_to_review,
    move_to_failed,
    move_to_archive,
    save_manifest,
    get_worker_id,
)

logger = structlog.get_logger(__name__)

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".tif"}
PDF_EXTENSIONS = {".pdf"}
SUPPORTED_EXTENSIONS = IMAGE_EXTENSIONS | PDF_EXTENSIONS | {".txt", ".md", ".docx"}


def _build_checkpointer():
    """MemorySaver by default (stateless design: review resume re-invokes the
    graph from the manifest, so no checkpoint persistence is needed — this also
    kills the unbounded per-doc checkpoint growth of SqliteSaver).

    Set MAILROOM_CHECKPOINTER=sqlite to opt back into the on-disk checkpointer
    (debugging/resume-across-restart experiments only).
    """
    import os

    if os.environ.get("MAILROOM_CHECKPOINTER", "memory") != "sqlite":
        logger.info("checkpointer_initialized", backend="memory")
        return MemorySaver()
    try:
        import sqlite3
        from langgraph.checkpoint.sqlite import SqliteSaver
        from pipeline.bins import get_base_dir

        db_path = get_base_dir() / "checkpoints.db"
        db_path.parent.mkdir(parents=True, exist_ok=True)
        conn = sqlite3.connect(str(db_path), check_same_thread=False)
        checkpointer = SqliteSaver(conn)
        try:
            checkpointer.setup()
        except Exception:
            pass
        logger.info("checkpointer_initialized", backend="sqlite", path=str(db_path))
        return checkpointer
    except Exception:
        logger.warning("sqlite_checkpointer_unavailable", fallback="memory")
    return MemorySaver()


def _ensure_dirs():
    ensure_dirs(
        inbox_dir(),
        processing_dir(),
        classified_dir(),
        review_dir(),
        failed_dir(),
        archive_dir(),
        manifests_dir(),
    )


def _read_file_text(file_path: Path) -> tuple[str, bool]:
    ext = file_path.suffix.lower()
    if ext in IMAGE_EXTENSIONS:
        return _extract_text_from_image(file_path)
    if ext in PDF_EXTENSIONS:
        return _extract_text_from_pdf(file_path)
    if ext == ".docx":
        return _extract_text_from_docx(file_path)
    try:
        text = file_path.read_text(errors="replace")
        if not text.strip():
            return ("", False)
        return (text, True)
    except Exception:
        try:
            text = file_path.read_bytes().decode("utf-8", errors="replace")
            return (text, bool(text.strip()))
        except Exception:
            return (f"[Unreadable file: {file_path.name}]", False)


def _render_doc_pages(file_path: Path) -> list[str]:
    """Render an input document to page-image data-URIs for vision-capable
    agents (PDFs page-by-page, image files passed through). Empty when vision is
    disabled or rendering is unavailable — text-only behaviour is unchanged."""
    try:
        from llm.vision import render_document_pages

        return render_document_pages(file_path)
    except Exception:
        logger.exception("doc_page_render_failed", file=str(file_path))
        return []


def _extract_text_from_docx(file_path: Path) -> tuple[str, bool]:
    """Extract text from .docx (paragraphs + tables) via python-docx.

    Previously .docx files fell through to the generic reader and were decoded
    as UTF-8 — i.e. zip binary garbage — which the classifier then tried to
    label. Unreadable files return the standard unreadable marker with
    ok=False so the pipeline routes them to review instead of misclassifying.
    """
    try:
        from docx import Document

        doc = Document(str(file_path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        if text.strip():
            logger.info("docx_text_extracted", file=file_path.name, chars=len(text))
            return (text, True)
        return ("", False)
    except ImportError:
        logger.warning("python_docx_missing", file=str(file_path))
        return (f"[Unreadable file: {file_path.name} — .docx support requires python-docx]", False)
    except Exception:
        logger.exception("docx_extraction_failed", file=str(file_path))
        return (f"[Unreadable file: {file_path.name} — not a valid .docx]", False)


def _extract_text_from_image(file_path: Path) -> tuple[str, bool]:
    logger.info("image_detected", file=str(file_path))
    try:
        from agents.image_extractor import ImageExtractor
        from observability.tracing import observation

        extractor = ImageExtractor()
        with observation("extract-image-text", input={"file": file_path.name}) as span:
            result = extractor.extract(file_path)
            if span is not None:
                span.update(output={"chars": len(result.get("text", ""))})
        text = result.get("text", "")
        if text:
            logger.info("image_extracted", file=file_path.name, chars=len(text))
            return (text, True)
    except Exception:
        logger.exception("image_extraction_failed", file=str(file_path))
    return (f"[Image file: {file_path.name} — text extraction failed]", False)


def _extract_text_from_pdf(file_path: Path) -> tuple[str, bool]:
    logger.info("pdf_detected", file=str(file_path))
    try:
        from agents.pdf_transcriber import PDFTranscriber
        from observability.tracing import observation

        transcriber = PDFTranscriber()
        with observation("transcribe-pdf", input={"file": file_path.name}) as span:
            result = transcriber.transcribe(file_path)
            if span is not None:
                span.update(
                    output={
                        "chars": len(result.get("markdown", "") or result.get("text", "")),
                        "method": result.get("method"),
                        "confidence": result.get("confidence"),
                    }
                )
        text = result.get("markdown", "") or result.get("text", "")
        if text:
            logger.info("pdf_transcribed", file=file_path.name, chars=len(text))
            return (text, True)
    except Exception:
        logger.exception("pdf_transcription_failed", file=str(file_path))
    try:
        import subprocess, tempfile
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            pass
        subprocess.run(["pdftotext", str(file_path), tmp.name], capture_output=True, timeout=30)
        text = Path(tmp.name).read_text(errors="replace")
        Path(tmp.name).unlink(missing_ok=True)
        if text.strip():
            logger.info("pdf_fallback_text", chars=len(text))
            return (text, True)
    except Exception:
        logger.exception("pdf_fallback_failed")
    return (f"[PDF file: {file_path.name} — transcription failed]", False)


def entry_route(state: dict) -> str:
    """Entry router: a review-resume re-invocation starts at fresh extraction
    (doc_type already known from the manifest); everything else goes through
    normal ingest → classify.

    The `review_decision == "approved"` guard is deliberate: only the
    resume-from-review path sets it, so a crashed/partial run can never be
    mistaken for a resume and skip classification (pilot: correspondence_01
    ended with output=null when a degraded second run took the extract branch
    without a real classification).
    """
    if (
        state.get("resume_extraction")
        and state.get("review_decision") == "approved"
        and state.get("doc_type")
    ):
        return "extract"
    return "ingest"


def _build_handoff_context(state: DocumentState) -> str | None:
    """Chained-eval handoff: prefix the sorter's classification (doc class +
    contract subtype) to the specialist's extraction call so it extracts with
    the expected field/clause set in mind (mirrors the sister repo's
    run_chained_eval pattern)."""
    doc_type = state.get("doc_type")
    if not doc_type:
        return None
    context = f"Sorter classification: doc_type={doc_type}"
    contract_subtype = state.get("contract_subtype")
    if doc_type == "contract" and contract_subtype:
        context += f" contract_subtype={contract_subtype}"
    confidence = state.get("classification_confidence")
    if confidence is not None:
        context += f" confidence={float(confidence):.2f}"
    context += ". Extract this document's fields accordingly, ensuring every expected item of this document class is captured."
    return context


def _build_specialist_dispatch():
    from pipeline.config import load_config
    cfg = load_config()
    doc_classes = cfg.get("doc_classes", [])
    dispatch = {}
    for cls in doc_classes:
        key = cls["key"]
        spec_name = cls.get("specialist", "")
        if spec_name == "contracts_specialist":
            dispatch[key] = _extract_contracts
        elif spec_name == "corporate_records_specialist":
            dispatch[key] = _extract_corporate_records
        elif spec_name == "due_diligence_specialist":
            dispatch[key] = _extract_due_diligence
        elif spec_name == "correspondence_specialist":
            dispatch[key] = _extract_correspondence
        elif spec_name == "compliance_specialist":
            dispatch[key] = _extract_compliance
        elif spec_name == "court_opinions_specialist":
            dispatch[key] = _extract_court_opinions
    return dispatch


def ingest_node(state: DocumentState) -> dict[str, Any]:
    _ensure_dirs()
    worker_id = get_worker_id()

    if state.get("file_path"):
        file_path = Path(state["file_path"])
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
    else:
        inbox = inbox_dir()
        files = list(inbox.glob("*"))
        if not files:
            raise RuntimeError("No files in inbox")
        file_path = files[0]

    if str(inbox_dir()) in str(file_path):
        file_path = claim_file(file_path, worker_id)

    doc_text, text_ok = _read_file_text(file_path)
    doc_pages = _render_doc_pages(file_path)

    matter_id = state.get("matter_id", "DEFAULT")
    manifest = DocumentManifest(
        matter_id=matter_id,
        original_filename=file_path.name,
        stage=PipelineStage.PROCESSING,
        trace_id=state.get("trace_id"),
    )
    manifest.touch()
    save_manifest(manifest)

    logger.info(
        "ingest",
        doc_id=manifest.doc_id,
        file=file_path.name,
        chars=len(doc_text),
        suffix=file_path.suffix,
        vision_pages=len(doc_pages),
    )
    # Write the processing-stage catalog record immediately so a crashed run is
    # visible to stuck-doc detection (`get_stuck_documents`) and `/ops/status`
    # instead of disappearing from the conveyor entirely.
    _catalog_upsert(
        {
            "doc_id": manifest.doc_id,
            "matter_id": matter_id,
            "original_filename": file_path.name,
            "doc_type": None,
            "stage": PipelineStage.PROCESSING.value,
        },
        stage=PipelineStage.PROCESSING.value,
    )
    return {
        "doc_id": manifest.doc_id,
        "matter_id": matter_id,
        "original_filename": file_path.name,
        "stage": PipelineStage.PROCESSING.value,
        "file_path": str(file_path),
        "doc_text": doc_text,
        "doc_pages": doc_pages,
        "classification_attempts": 0,
        "extraction_attempts": 0,
        "retry_count": 0,
        "conflict_detected": False,
        "error_message": None if text_ok else f"Could not extract text from {file_path.suffix} file",
    }


def classify_node(state: DocumentState) -> dict[str, Any]:
    doc_text = state.get("doc_text", "")
    if not doc_text or not doc_text.strip():
        logger.warning("empty_doc_text_classify", doc_id=state.get("doc_id"))
        # Empty/unreadable text can never be classified — route straight to
        # human review instead of burning a retry call on the same empty text
        # (which also clobbered this escalation reason on the retry). Setting
        # classification_attempts past retry_max makes after_classify send it
        # to review immediately.
        retry_max = get_confidence_thresholds().get("retry_max", 1)
        return {
            "doc_type": "correspondence",
            "classification_confidence": 0.1,
            "classification_attempts": max(state.get("classification_attempts", 0), retry_max) + 1,
            "stage": PipelineStage.CLASSIFIED.value,
            "escalation_reason": "Empty or unreadable document content",
            "transient_error": False,
        }

    from agents.sorter import SorterAgent
    from llm.retry import is_transient_error
    from pipeline.guards import guard_classification

    sorter = SorterAgent()
    attempts = state.get("classification_attempts", 0)
    try:
        # Vendored LangChain sorter returns the 4-tuple (doc_type,
        # contract_subtype, confidence, reasoning).
        doc_type, contract_subtype, confidence, reasoning = sorter.classify(
            doc_text, pages=state.get("doc_pages")
        )
    except Exception as exc:
        if is_transient_error(exc):
            # Provider-side transient failure (connection/timeout/rate-limit/
            # 5xx). Do NOT increment the confidence retry budget; routing
            # retries this same node via the `classify` self-loop.
            transient = state.get("transient_retries", 0) + 1
            logger.warning(
                "classify_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries": transient,
                "classification_attempts": attempts,
                "stage": PipelineStage.CLASSIFIED.value,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during classification",
            }
        raise
    attempts = attempts + 1

    guard = guard_classification(
        {
            "doc_type": doc_type,
            "classification_confidence": confidence,
            "contract_subtype": contract_subtype,
        }
    )
    if not guard["ok"]:
        # Never trust out-of-range confidence or an invalid contract subtype;
        # leave doc_type untouched so routing's unknown-type check sends it to
        # human review. Record the guardrail on state so it is scored
        # (guardrail_triggered) and visible in traces.
        logger.warning("classification_guardrail_triggered", doc_id=state.get("doc_id"), issues=guard["issues"])
        confidence = guard.get("confidence", 0.1)

    logger.info(
        "classified",
        doc_type=doc_type,
        contract_subtype=contract_subtype,
        confidence=confidence,
        attempts=attempts,
    )
    return {
        "doc_type": doc_type,
        "contract_subtype": contract_subtype,
        "classification_confidence": confidence,
        "classification_attempts": attempts,
        "classification_guardrail": guard["issues"],
        "stage": PipelineStage.CLASSIFIED.value,
        "escalation_reason": reasoning
        if confidence < get_confidence_thresholds().get("high", 0.95)
        else None,
        "transient_error": False,
    }


def retry_classify_node(state: DocumentState) -> dict[str, Any]:
    from agents.sorter import SorterAgent
    from llm.retry import is_transient_error

    sorter = SorterAgent()
    doc_text = state.get("doc_text", "")
    attempts = state.get("classification_attempts", 0)

    prev_type = state.get("doc_type", "")
    prev_confidence = state.get("classification_confidence", 0)
    augmented_text = (
        f"RE-EVALUATION REQUESTED - previous classification was '{prev_type}' with "
        f"confidence {prev_confidence:.2f}. Please re-examine this document independently:\n\n"
        f"{doc_text[:12000]}"
    )
    try:
        doc_type, contract_subtype, confidence, reasoning = sorter.classify(
            augmented_text, pages=state.get("doc_pages")
        )
    except Exception as exc:
        if is_transient_error(exc):
            transient = state.get("transient_retries", 0) + 1
            logger.warning(
                "retry_classify_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries": transient,
                "classification_attempts": attempts,
                "stage": PipelineStage.CLASSIFIED.value,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during re-classification",
            }
        raise
    attempts = attempts + 1

    # Same deterministic guard as the first-pass classify: never trust an
    # out-of-range confidence, unknown doc type, or invalid contract subtype
    # from the retry either (it would otherwise route to extract directly).
    from pipeline.guards import guard_classification

    guard = guard_classification(
        {
            "doc_type": doc_type,
            "classification_confidence": confidence,
            "contract_subtype": contract_subtype,
        }
    )
    if not guard["ok"]:
        logger.warning(
            "retry_classification_guardrail_triggered",
            doc_id=state.get("doc_id"),
            issues=guard["issues"],
        )
        confidence = guard.get("confidence", 0.1)

    logger.info(
        "retry_classified",
        doc_type=doc_type,
        contract_subtype=contract_subtype,
        confidence=confidence,
        attempts=attempts,
    )
    return {
        "doc_type": doc_type,
        "contract_subtype": contract_subtype,
        "classification_confidence": confidence,
        "classification_attempts": attempts,
        "retry_count": state.get("retry_count", 0) + 1,
        "classification_guardrail": guard["issues"],
        "stage": PipelineStage.CLASSIFIED.value,
        "escalation_reason": reasoning
        if confidence < get_confidence_thresholds().get("high", 0.95)
        else None,
        "transient_error": False,
    }


def _fetch_matter_context(state: dict) -> list[dict]:
    """Best-effort fetch of archived matter records for the Boss / conflict
    detection. Returns a list of {doc_id, doc_type, extracted_data} dicts;
    never raises (DB unavailable → empty context)."""
    matter_id = state.get("matter_id")
    doc_id = state.get("doc_id")
    if not matter_id:
        return []
    try:
        from storage.catalog import get_matter_documents

        rows = _run_coro(lambda: get_matter_documents(matter_id))
        return [
            {
                "doc_id": r.doc_id,
                "doc_type": r.doc_type,
                "extracted_data": r.extracted_data or {},
                "stage": r.stage,
            }
            for r in rows
            if r.doc_id != doc_id and r.stage == "archived"
        ]
    except Exception:
        logger.exception("matter_context_fetch_failed")
        return []


def _normalized_compare(a, b) -> bool:
    """Normalized value comparison for conflict detection (list-aware)."""
    if isinstance(a, list) and isinstance(b, list):
        na = {_norm_str(x) for x in a}
        nb = {_norm_str(x) for x in b}
        return na == nb
    return _norm_str(a) == _norm_str(b)


def _norm_str(v) -> str:
    try:
        from observability.field_scoring import normalize_text

        return normalize_text(v)
    except Exception:
        return str(v).strip().lower()


def _detect_conflict(state: dict, extracted_data: dict | None) -> tuple[bool, list[str]]:
    """Deterministically compare a fresh extraction against archived records
    of the same matter. A conflict exists when the same field is populated on
    both sides with a different normalized value (e.g. two contracts in one
    matter claiming different governing laws or parties).

    Returns (conflict_detected, details). Best-effort: no DB → no conflict.
    """
    if not extracted_data:
        return False, []
    # Only fields present in the schema are conflict-relevant (ignore pipeline
    # metadata keys like `_report`).
    schema_fields = set()
    try:
        from schemas.documents import get_extraction_schema

        model = get_extraction_schema(state.get("doc_type") or "")
        if model is not None:
            schema_fields = set(model.model_fields.keys())
    except Exception:
        pass

    details: list[str] = []
    for record in _fetch_matter_context(state):
        prior = record.get("extracted_data") or {}
        if not prior:
            continue
        for field in sorted(schema_fields):
            new_val = extracted_data.get(field)
            old_val = prior.get(field)
            if new_val is None or old_val is None:
                continue
            if isinstance(new_val, str) and not new_val.strip():
                continue
            if isinstance(old_val, str) and not old_val.strip():
                continue
            if isinstance(new_val, (list, dict)) and not new_val:
                continue
            if isinstance(old_val, (list, dict)) and not old_val:
                continue
            if _normalized_compare(new_val, old_val):
                continue
            details.append(
                f"field '{field}' differs from archived record "
                f"{record.get('doc_id', '?')}: '{old_val}' vs '{new_val}'"
            )
            break  # one conflict per prior record is enough to escalate
    return bool(details), details


def extract_node(state: DocumentState) -> dict[str, Any]:
    doc_type = state.get("doc_type", "")
    doc_text = state.get("doc_text", "")
    doc_pages = state.get("doc_pages") or []

    dispatch = _build_specialist_dispatch()
    extractor = dispatch.get(
        doc_type,
        lambda t, pages=None, handoff_context=None: {"confidence": 0.3, "_unsupported": True},
    )
    handoff_context = _build_handoff_context(state)
    attempts = state.get("extraction_attempts", 0)
    try:
        result = extractor(doc_text, doc_pages, handoff_context)
    except Exception as exc:
        from llm.retry import is_transient_error

        if is_transient_error(exc):
            # Transient provider failure: retry the same node without burning
            # the extraction retry budget (routed via the `extract` self-loop).
            transient = state.get("transient_retries", 0) + 1
            logger.warning(
                "extract_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries": transient,
                "extraction_attempts": attempts,
                "extraction_confidence": 0.0,
                "extracted_data": None,
                "stage": PipelineStage.CLASSIFIED.value,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during extraction",
            }
        # Non-transient exception: convert it into a parse-level failure so the
        # deterministic guardrail clamps confidence and routing sends the doc
        # to retry → review instead of crashing the run silently.
        logger.exception(
            "extraction_exception",
            doc_id=state.get("doc_id"),
            doc_type=doc_type,
            error=str(exc)[:300],
        )
        result = {"_parse_error": True, "_exception": str(exc), "confidence": 0.0}
    confidence = result.pop("confidence", None)
    attempts = attempts + 1

    from pipeline.guards import apply_extraction_guard

    guard, confidence = apply_extraction_guard(doc_type, result, confidence, attempts=attempts)

    # Deterministic conflict detection against archived matter records: a
    # materially different value for the same schema field (governing law,
    # parties, effective date, amounts) is escalated to the Boss for
    # adjudication instead of silently overwriting matter history.
    conflict_detected, conflict_details = _detect_conflict(state, result)

    logger.info(
        "extracted",
        doc_type=doc_type,
        confidence=confidence,
        attempts=attempts,
        conflict_detected=conflict_detected,
    )
    return {
        "extracted_data": result,
        "extraction_confidence": confidence,
        "extraction_attempts": attempts,
        "extraction_guardrail": guard["issues"],
        "conflict_detected": conflict_detected,
        "conflict_details": conflict_details,
        "escalation_reason": "; ".join(conflict_details)
        if conflict_detected
        else state.get("escalation_reason"),
        "transient_error": False,
    }


def _extract_contracts(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.contracts_specialist import ContractsSpecialist
    return ContractsSpecialist(handoff_context=handoff_context).extract(doc_text, pages=pages)


def _extract_corporate_records(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.corporate_records_specialist import CorporateRecordsSpecialist
    return CorporateRecordsSpecialist().extract(doc_text, pages=pages, handoff_context=handoff_context)


def _extract_due_diligence(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.due_diligence_specialist import DueDiligenceSpecialist
    return DueDiligenceSpecialist().extract(doc_text, pages=pages, handoff_context=handoff_context)


def _extract_correspondence(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.correspondence_specialist import CorrespondenceSpecialist
    return CorrespondenceSpecialist().extract(doc_text, pages=pages, handoff_context=handoff_context)


def _extract_compliance(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.compliance_specialist import ComplianceSpecialist
    return ComplianceSpecialist().extract(doc_text, pages=pages, handoff_context=handoff_context)


def _extract_court_opinions(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.court_opinions_specialist import CourtOpinionsSpecialist
    return CourtOpinionsSpecialist().extract(doc_text, pages=pages, handoff_context=handoff_context)


def retry_extract_node(state: DocumentState) -> dict[str, Any]:
    doc_type = state.get("doc_type", "")
    doc_text = state.get("doc_text", "")
    doc_pages = state.get("doc_pages") or []
    prev_extracted = state.get("extracted_data", {})
    attempts = state.get("extraction_attempts", 0)

    augmented_text = (
        f"RE-EXTRACTION REQUESTED - previous extraction was low-confidence. "
        f"Please re-examine this document independently. Previous attempt found: {prev_extracted}\n\n"
        f"{doc_text[:25000]}"
    )

    dispatch = _build_specialist_dispatch()
    extractor = dispatch.get(
        doc_type,
        lambda t, pages=None, handoff_context=None: {"confidence": 0.3, "_unsupported": True},
    )
    handoff_context = _build_handoff_context(state)
    try:
        result = extractor(augmented_text, doc_pages, handoff_context)
    except Exception as exc:
        from llm.retry import is_transient_error

        if is_transient_error(exc):
            transient = state.get("transient_retries", 0) + 1
            logger.warning(
                "retry_extract_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries": transient,
                "extraction_attempts": attempts,
                "extraction_confidence": 0.0,
                "extracted_data": None,
                "stage": PipelineStage.CLASSIFIED.value,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during re-extraction",
            }
        logger.exception(
            "retry_extraction_exception",
            doc_id=state.get("doc_id"),
            doc_type=doc_type,
            error=str(exc)[:300],
        )
        result = {"_parse_error": True, "_exception": str(exc), "confidence": 0.0}
    confidence = result.pop("confidence", None)
    attempts = attempts + 1

    from pipeline.guards import apply_extraction_guard

    guard, confidence = apply_extraction_guard(doc_type, result, confidence, attempts=attempts)

    conflict_detected, conflict_details = _detect_conflict(state, result)

    logger.info(
        "retry_extracted",
        doc_type=doc_type,
        confidence=confidence,
        attempts=attempts,
        conflict_detected=conflict_detected,
    )
    return {
        "extracted_data": result,
        "extraction_confidence": confidence,
        "extraction_attempts": attempts,
        "retry_count": state.get("retry_count", 0) + 1,
        "extraction_guardrail": guard["issues"],
        "conflict_detected": conflict_detected,
        "conflict_details": conflict_details,
        "escalation_reason": "; ".join(conflict_details)
        if conflict_detected
        else state.get("escalation_reason"),
        "transient_error": False,
    }


def human_review_node(state: DocumentState) -> dict[str, Any]:
    file_path_str = state.get("file_path", "")
    esc_reason = state.get("escalation_reason", "Unknown reason")
    doc_id = state.get("doc_id", "")

    logger.info("human_review_required", doc_id=doc_id, reason=esc_reason)

    if file_path_str:
        manifest = DocumentManifest(
            doc_id=doc_id,
            matter_id=state.get("matter_id", "DEFAULT"),
            original_filename=state.get("original_filename", ""),
            stage=PipelineStage.REVIEW,
            doc_type=state.get("doc_type"),
            contract_subtype=state.get("contract_subtype"),
            classification_confidence=state.get("classification_confidence"),
            extracted_data=state.get("extracted_data"),
            extraction_confidence=state.get("extraction_confidence"),
            escalation_reason=esc_reason,
            trace_id=state.get("trace_id"),
            classification_attempts=state.get("classification_attempts", 0),
            extraction_attempts=state.get("extraction_attempts", 0),
        )
        move_to_review(Path(file_path_str), manifest)
        # Persist the review position in the catalog so `/ops/status`
        # review_queue and error-rate stats see it (they query the catalog).
        _catalog_upsert(
            {
                "doc_id": doc_id,
                "matter_id": state.get("matter_id", "DEFAULT"),
                "original_filename": state.get("original_filename", ""),
                "doc_type": state.get("doc_type"),
                "contract_subtype": state.get("contract_subtype"),
                "classification_confidence": state.get("classification_confidence"),
                "extraction_confidence": state.get("extraction_confidence"),
                "extracted_data": state.get("extracted_data"),
                "escalation_reason": esc_reason,
                "trace_id": state.get("trace_id"),
                "stage": PipelineStage.REVIEW.value,
            },
            stage=PipelineStage.REVIEW.value,
        )

    return {
        "stage": PipelineStage.REVIEW.value,
        "escalation_reason": esc_reason,
        # Sentinal for graph termination, NOT a human decision: the document
        # was routed to review and is awaiting a human. "rejected" would leak
        # into traces/manifests as a verdict the human never gave.
        "review_decision": "pending_review",
    }


def boss_escalation_node(state: DocumentState) -> dict[str, Any]:
    from agents.boss import BossAgent

    boss = BossAgent()
    manifest_data = {
        "doc_id": state.get("doc_id"),
        "doc_type": state.get("doc_type"),
        "classification_confidence": state.get("classification_confidence"),
        "extraction_confidence": state.get("extraction_confidence"),
        "extracted_data": state.get("extracted_data"),
        "escalation_reason": state.get("escalation_reason"),
    }

    # Give the Boss the archived matter context it adjudicates against, so its
    # decision is grounded in the actual conflicting records (it receives only
    # the doc manifest otherwise). Best-effort: empty context when unavailable.
    matter_context = _fetch_matter_context(state)

    result = boss.adjudicate(manifest_data, matter_context=matter_context)
    decision = result.get("decision", "review")
    reasoning = result.get("reasoning", "")

    logger.info(
        "boss_decision",
        doc_id=state.get("doc_id"),
        decision=decision,
        context_records=len(matter_context),
    )
    return {
        "review_decision": decision,
        "escalation_reason": f"Boss: {reasoning}",
    }


def compile_report_node(state: DocumentState) -> dict[str, Any]:
    from llm.client import get_llm

    manifest_data = {
        "doc_id": state.get("doc_id"),
        "matter_id": state.get("matter_id"),
        "doc_type": state.get("doc_type"),
        "contract_subtype": state.get("contract_subtype"),
        "classification_confidence": state.get("classification_confidence"),
        "extraction_confidence": state.get("extraction_confidence"),
        "extracted_data": state.get("extracted_data"),
    }

    llm_client, model = get_llm("reporter")
    from agents.reporter import compile_matter_record
    report = compile_matter_record(manifest_data, llm_client, model)

    logger.info("report_compiled", doc_id=state.get("doc_id"))
    return {
        "extracted_data": {
            **state.get("extracted_data", {}),
            "_report": report,
        },
    }


def _catalog_upsert(state: dict, *, stage: str | None = None, update_only: bool = False) -> None:
    """Best-effort write/update the catalog record for a document state.

    The catalog is the durable conveyor position: every stage transition that
    moves a document (processing → review/failed/archived) upserts the row so
    `/ops/status`, stuck-doc detection, matter listings, and error rates reflect
    reality even when the run ends in review or fails. Never raises.
    """
    try:
        from schemas.matter import Matter
        from storage.catalog import write_document_record as _write_doc, write_matter_record as _write_matter

        doc_id = state.get("doc_id") or ""
        matter_id = state.get("matter_id") or "DEFAULT"
        if not doc_id:
            return

        doc_record = {
            "doc_id": doc_id,
            "matter_id": matter_id,
            "original_filename": state.get("original_filename", ""),
            "doc_type": state.get("doc_type", "unknown"),
            "contract_subtype": state.get("contract_subtype"),
            "stage": stage or state.get("stage", "cataloged"),
            "classification_confidence": state.get("classification_confidence"),
            "extraction_confidence": state.get("extraction_confidence"),
            "extracted_data": state.get("extracted_data"),
            "escalation_reason": state.get("escalation_reason") or state.get("error_message"),
            "trace_id": state.get("trace_id"),
        }

        def _sync_write():
            async def _runner():
                if not update_only:
                    matter = Matter(
                        matter_id=matter_id,
                        name=matter_id,
                        client_name="auto-created",
                        practice_area="transactional",
                    )
                    await _write_matter(matter)
                await _write_doc(doc_record)

            _run_coro(_runner)

        _sync_write()
        logger.debug("catalog_upserted", doc_id=doc_id, stage=doc_record["stage"])
    except Exception:
        logger.exception("catalog_upsert_error")


def catalog_write_node(state: DocumentState) -> dict[str, Any]:
    doc_id = state.get("doc_id", "")
    logger.info("catalog_write", doc_id=doc_id)
    # The catalog_write node runs BEFORE archive: the doc_type/confidence/
    # extraction fields are final here, but the terminal stage is not yet known
    # (archive_node completes the move). Write with the current stage; the
    # archive node upserts stage=archived afterwards so the catalog never
    # permanently shows a document as merely "classified".
    _catalog_upsert(state, stage=state.get("stage", "classified"))
    return {}


def archive_node(state: DocumentState) -> dict[str, Any]:
    manifest = DocumentManifest(
        doc_id=state.get("doc_id", ""),
        matter_id=state.get("matter_id", "DEFAULT"),
        original_filename=state.get("original_filename", ""),
        stage=PipelineStage.ARCHIVED,
        doc_type=state.get("doc_type", "unknown"),
        contract_subtype=state.get("contract_subtype"),
        classification_confidence=state.get("classification_confidence"),
        extracted_data=state.get("extracted_data"),
        extraction_confidence=state.get("extraction_confidence"),
        trace_id=state.get("trace_id"),
        escalation_reason=state.get("escalation_reason"),
        review_decision=state.get("review_decision"),
        classification_attempts=state.get("classification_attempts", 0),
        extraction_attempts=state.get("extraction_attempts", 0),
    )

    file_path_str = state.get("file_path", "")
    if not file_path_str:
        logger.error("archive_no_file_path", doc_id=manifest.doc_id)
        return {"stage": PipelineStage.FAILED.value, "error_message": "No file path in state"}

    file_path = Path(file_path_str)
    if not file_path.exists():
        processing_root = processing_dir()
        candidates = list(processing_root.rglob(state.get("original_filename", "*.txt")))
        if candidates:
            file_path = candidates[0]
        else:
            logger.error("archive_file_not_found", doc_id=manifest.doc_id, path=file_path_str)
            return {"stage": PipelineStage.FAILED.value, "error_message": f"File not found: {file_path_str}"}

    from agents.archivist import archive_document
    # Hash-chained audit trail: the new entry must link to the previous entry
    # for THIS doc_id (review decisions, re-runs, resumes all append to the
    # same chain). A stale `""` would break `verify_chain` for any second
    # event on the same document.
    prev_audit_hash = _latest_audit_hash(manifest.doc_id)
    archive_path, audit_entry = archive_document(manifest, file_path, prev_audit_hash=prev_audit_hash)

    _write_audit_log(audit_entry)

    # Final conveyor position: the catalog record (created at ingest/catalog
    # write) must show archived, not classified — archive is the terminal stage.
    _catalog_upsert(
        {
            "doc_id": manifest.doc_id,
            "matter_id": manifest.matter_id,
            "original_filename": manifest.original_filename,
            "doc_type": manifest.doc_type,
            "contract_subtype": manifest.contract_subtype,
            "classification_confidence": manifest.classification_confidence,
            "extraction_confidence": manifest.extraction_confidence,
            "extracted_data": manifest.extracted_data,
            "escalation_reason": manifest.escalation_reason,
            "trace_id": manifest.trace_id,
            "stage": PipelineStage.ARCHIVED.value,
        },
        stage=PipelineStage.ARCHIVED.value,
        update_only=True,
    )

    logger.info("pipeline_complete", doc_id=manifest.doc_id, archive=str(archive_path))
    return {"stage": PipelineStage.ARCHIVED.value}


def _run_coro(coro):
    """Run a coroutine from a sync context: schedule it on the running loop
    when one exists (thread-safe), otherwise run a fresh loop.

    `asyncio.get_event_loop()` is deprecated when no loop is running, and
    graph nodes execute both from the watcher's daemon threads (no loop) and
    from the API's `asyncio.to_thread` (a running loop in another thread), so
    both branches are needed.
    """
    import asyncio

    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(coro())
    import concurrent.futures

    future = asyncio.run_coroutine_threadsafe(coro(), loop)
    return future.result(timeout=10)


def _latest_audit_hash(doc_id: str) -> str:
    """Best-effort fetch of the last entry_hash for this doc_id (the previous
    link of the hash chain). Returns "" when no entries exist yet (a fresh
    chain) or the DB is unavailable — a broken link can never be caused by
    this: `build_audit_entry` uses whatever we return as prev_hash, and
    `verify_chain` recomputes from timestamps."""
    if not doc_id:
        return ""
    try:
        from storage.audit_log import get_latest_audit_hash

        return _run_coro(lambda: get_latest_audit_hash(doc_id))
    except Exception:
        logger.exception("latest_audit_hash_fetch_failed", doc_id=doc_id)
        return ""


def _write_audit_log(entry):
    try:
        import asyncio
        from storage.audit_log import write_audit_entry

        async def _write():
            await write_audit_entry(entry)

        _run_coro(_write)
    except Exception:
        logger.exception("audit_log_write_error")


def _persist_scores(state: dict, scores: dict):
    if not scores or not state.get("doc_id"):
        return
    try:
        from storage.catalog import update_document_scores

        async def _write():
            await update_document_scores(state["doc_id"], scores)

        _run_coro(_write)
    except Exception:
        logger.exception("scores_persist_error")


def _bounded(fn):
    """Node wrapper enforcing the per-run hard cutoff: wall-clock deadline and
    cumulative output-token budget. Raises RunDeadlineExceeded /
    RunBudgetExceeded, which `_execute_run` catches and finalizes as an
    aborted run (failed bin + scores)."""
    from pipeline.limits import check_run_deadline, check_token_budget

    @functools.wraps(fn)
    def wrapper(state):
        check_run_deadline(state.get("run_deadline"))
        check_token_budget()
        return fn(state)

    return wrapper


def build_graph(checkpointer=None):
    if checkpointer is None:
        checkpointer = _build_checkpointer()

    workflow = StateGraph(DocumentState)

    # Node names stay stable (best practice); per-run values go in metadata.
    # Every node is bounded: the run deadline and token budget are enforced at
    # each boundary so a stuck run is cut off as soon as its budget is spent.
    workflow.add_node("ingest", traced_node("ingest-document")(_bounded(ingest_node)))
    workflow.add_node("classify", traced_node("classify-document")(_bounded(classify_node)))
    workflow.add_node("retry_classify", traced_node("classify-document")(_bounded(retry_classify_node)))
    workflow.add_node("extract", traced_node("extract-fields")(_bounded(extract_node)))
    workflow.add_node("retry_extract", traced_node("extract-fields")(_bounded(retry_extract_node)))
    workflow.add_node("human_review", traced_node("route-for-review")(_bounded(human_review_node)))
    workflow.add_node("boss_escalation", traced_node("adjudicate-conflict")(_bounded(boss_escalation_node)))
    workflow.add_node("compile_report", traced_node("compile-report")(_bounded(compile_report_node)))
    workflow.add_node("catalog_write", traced_node("write-catalog")(_bounded(catalog_write_node)))
    workflow.add_node("archive", traced_node("archive-document")(_bounded(archive_node)))

    workflow.add_conditional_edges(START, entry_route, {
        "ingest": "ingest",
        "extract": "extract",
    })
    workflow.add_edge("ingest", "classify")

    workflow.add_conditional_edges("classify", after_classify, {
        "classify": "classify",  # transient-error self-loop (same node, LLM-level retry)
        "retry_classify": "retry_classify",
        "extract": "extract",
        "human_review": "human_review",
    })

    workflow.add_conditional_edges("retry_classify", after_retry_classify, {
        "classify": "classify",
        "extract": "extract",
        "human_review": "human_review",
    })

    workflow.add_conditional_edges("extract", after_extraction, {
        "extract": "extract",  # transient-error self-loop (same node, LLM-level retry)
        "retry_extract": "retry_extract",
        "compile_report": "compile_report",
        "human_review": "human_review",
        "boss_escalation": "boss_escalation",
    })

    workflow.add_conditional_edges("retry_extract", after_retry_extraction, {
        "extract": "extract",
        "compile_report": "compile_report",
        "human_review": "human_review",
        "boss_escalation": "boss_escalation",
    })

    workflow.add_conditional_edges("boss_escalation", after_boss, {
        "compile_report": "compile_report",
        "human_review": "human_review",
    })

    workflow.add_conditional_edges("human_review", after_human_review, {
        "compile_report": "compile_report",
        "failed": END,
    })

    workflow.add_edge("compile_report", "catalog_write")
    workflow.add_edge("catalog_write", "archive")
    workflow.add_edge("archive", END)

    return workflow.compile(checkpointer=checkpointer)


def _existing_processing_doc_id(original_filename: str) -> str | None:
    """Find the doc_id of an in-flight manifest for this filename.

    A run that crashed after ingest already saved a processing-stage manifest
    (and a catalog row); the abort path must reuse that doc_id so the failed
    manifest/catalog record supersede the same document instead of orphaning
    the ingest manifest and minting a second identity.
    """
    if not original_filename:
        return None
    try:
        import json as _json
        from pipeline.bins import manifests_dir

        mdir = manifests_dir()
        if not mdir.exists():
            return None
        for mf in mdir.glob("*.json"):
            try:
                data = _json.loads(mf.read_text())
            except Exception:
                continue
            if data.get("original_filename") != original_filename:
                continue
            if data.get("stage") == PipelineStage.PROCESSING.value:
                return data.get("doc_id")
    except Exception:
        logger.exception("abort_doc_id_lookup_failed")
    return None


def _finalize_aborted(initial_state: dict, reason: str) -> dict:
    """Turn a run that hit a hard limit (or crashed) into a failed result.

    Moves the file to the failed bin with a manifest noting the abort, and
    returns a result dict that still carries doc/attempt fields so the run is
    scored (run_aborted=1) and visible in the catalog instead of stranding in
    processing/.
    """
    from pipeline.bins import move_to_failed, save_manifest
    from schemas.manifest import DocumentManifest, PipelineStage

    state = dict(initial_state)
    # Reuse the ingest manifest's doc_id when the run crashed after ingest, so
    # the aborted manifest supersedes the processing manifest (same identity).
    # Passed explicitly — DocumentManifest would otherwise mint a fresh UUID.
    aborted_doc_id = state.get("doc_id") or _existing_processing_doc_id(
        state.get("original_filename", "")
    )
    manifest_kwargs = dict(
        matter_id=state.get("matter_id", "DEFAULT"),
        original_filename=state.get("original_filename", ""),
        stage=PipelineStage.FAILED,
        doc_type=state.get("doc_type"),
        contract_subtype=state.get("contract_subtype"),
        classification_confidence=state.get("classification_confidence"),
        classification_attempts=state.get("classification_attempts", 0),
        extracted_data=state.get("extracted_data"),
        extraction_confidence=state.get("extraction_confidence"),
        extraction_attempts=state.get("extraction_attempts", 0),
        escalation_reason=f"run aborted: {reason}",
        trace_id=state.get("trace_id"),
    )
    if aborted_doc_id:
        manifest_kwargs["doc_id"] = aborted_doc_id
    manifest = DocumentManifest(**manifest_kwargs)
    state["doc_id"] = manifest.doc_id
    state["stage"] = PipelineStage.FAILED.value
    state["run_aborted"] = True
    state["error_message"] = f"run aborted: {reason}"

    file_path_str = state.get("file_path") or ""
    if file_path_str:
        try:
            move_to_failed(Path(file_path_str))
        except Exception:
            logger.exception("abort_move_to_failed_error", file=file_path_str)
    try:
        save_manifest(manifest)
    except Exception:
        logger.exception("abort_manifest_save_error", doc_id=manifest.doc_id)
    try:
        _write_catalog_record(state)
    except Exception:
        logger.exception("abort_catalog_write_error", doc_id=manifest.doc_id)
    return state


def _write_catalog_record(state: dict):
    """Persist a minimal catalog record (used for aborted runs that never
    reach the catalog_write node, so they show up in compare_runs)."""
    import asyncio
    from storage.catalog import write_document_record

    doc_record = {
        "doc_id": state.get("doc_id", ""),
        "matter_id": state.get("matter_id", "DEFAULT"),
        "original_filename": state.get("original_filename", ""),
        "doc_type": state.get("doc_type", "unknown"),
        "stage": state.get("stage", "failed"),
        "classification_confidence": state.get("classification_confidence"),
        "extraction_confidence": state.get("extraction_confidence"),
        "extracted_data": state.get("extracted_data"),
        "escalation_reason": state.get("escalation_reason") or state.get("error_message"),
        "trace_id": state.get("trace_id"),
    }

    async def _write():
        await write_document_record(doc_record)

    _run_coro(_write)


# Cap for the judge-visible document text in the pipeline-result generation.
# Long contracts (e.g. MAUD merger agreements, ~800KB) would blow up the single
# cumulative judge call otherwise.
PIPELINE_RESULT_TEXT_LIMIT = 100_000


def _emit_pipeline_result(root, result: dict, state: dict, judge_required: bool | None = None) -> None:
    """Emit the single `pipeline-result` generation observation per document.

    This is the ONLY observation the live evaluation rule matches
    (`scripts/sync_evaluators.py`, `mailroom-pipeline-rule`), so exactly one
    cumulative LLM-as-a-Judge call scores each document — classification
    correctness + extraction correctness + completeness in one pass. Output is
    the curated pipeline result. The input depends on the run mode:

    - Grounded runs (ground_truth carries `expected_fields`, i.e. pilot runs):
      the judge input is a labeled, pretty-printed expected-fields block. The
      output contains the labeled pipeline result and extracted fields. The
      full document text is NOT sent, which cuts the per-doc judge input from
      up to 100k chars (~25k tokens) to ~1-3k chars. The expected values ARE
      the ground truth, so the document text adds no verification power.
    - Live runs (no ground truth): the judge gets the (truncated) document
      text so it can verify grounding by rubric alone.

    Judge gating (issues #4/#5): for grounded runs the deterministic
    field-type-aware scorer (`observability/field_scoring.py`) runs first.
    When its verdict is unambiguous — every scored field clearly correct
    (above the ambiguous band) OR clearly wrong (below it) — `judge_required`
    is False and the generation is NOT emitted, so neither evaluator rule
    fires and two judge LLM calls are saved per document. When any field lands
    in the ambiguous band (0.5-0.85), or there is no ground truth at all
    (`judge_required is None`, live runs), the generation is emitted as
    before.

    No-ops when tracing is disabled (root None).

    A run that ends in `review` is NOT a final pipeline outcome: the document
    is awaiting (or has already received) a human decision, and the resumed
    run (if approved) re-archives under the same deterministic trace id. The
    generation is therefore suppressed for review-routed runs — the resumed
    run emits the single authoritative `pipeline-result`, so the evaluator
    fires exactly once per document trace instead of twice (once judged MISS
    for the review stage, once CORRECT after resume).
    """
    if root is None:
        return
    if result.get("stage") in ("review", "failed"):
        logger.info("pipeline_result_suppressed_non_terminal_stage", stage=result.get("stage"))
        return
    if judge_required is False:
        logger.info("pipeline_result_suppressed_deterministic_verdict")
        return
    import json

    ground_truth = state.get("ground_truth")
    grounded = bool(ground_truth and ground_truth.get("expected_fields"))
    extracted_data = result.get("extracted_data") or {}
    # `_report` is a derived catalog summary and may contain a full recursive
    # copy of the extraction. It is not part of any specialist schema and must
    # never be sent to the evaluator.
    judge_extracted_data = {
        key: value for key, value in extracted_data.items() if not key.startswith("_")
    }
    output = {
        "stage": result.get("stage"),
        "doc_type": result.get("doc_type"),
        "classification_confidence": result.get("classification_confidence"),
        "extraction_confidence": result.get("extraction_confidence"),
        "extracted_data": judge_extracted_data,
        "escalation_reason": result.get("escalation_reason"),
        "review_decision": result.get("review_decision"),
        "run_aborted": bool(result.get("run_aborted")),
        "error_message": result.get("error_message"),
    }
    if grounded:
        # Skip the document text entirely. The expected fields are the only
        # judge input; extracted fields are in the output, avoiding duplication.
        gen_input = (
            "GROUNDED EVALUATION INPUT\n"
            "The following fields are the literal expected values for this document.\n"
            "Compare them only with output.extracted_data for this same document.\n\n"
            "EXPECTED_FIELDS\n"
            "```json\n"
            f"{json.dumps(ground_truth['expected_fields'], ensure_ascii=False, indent=2)}\n"
            "```"
        )
        metadata = {"pipeline": "mailroom", "grounded": True, "input_format": "expected-fields-only"}
    else:
        doc_text = result.get("doc_text") or state.get("doc_text") or ""
        gen_input = doc_text[:PIPELINE_RESULT_TEXT_LIMIT]
        metadata = {
            "pipeline": "mailroom",
            "truncated": len(doc_text) > PIPELINE_RESULT_TEXT_LIMIT,
        }
    if ground_truth:
        # When the caller knows the expected outcome (pilot runs pass the
        # manifest ground truth), expose it here so the live evaluator can
        # decide a CORRECT/PARTIAL/MISS verdict against the ACTUAL truth
        # instead of judging by rubric alone. Expected fields already live in
        # the labeled input block, so do not duplicate them in the output.
        output["ground_truth"] = {
            key: value for key, value in ground_truth.items() if key != "expected_fields"
        }
    with observation(
        "pipeline-result",
        as_type="generation",
        input=gen_input,
        metadata=metadata,
    ) as gen:
        if gen is not None:
            gen.update(output=output)


def _execute_run(
    initial_state: DocumentState,
    seed: str,
    trace_input: dict,
    attempt: int = 0,
    source: str | None = None,
    ground_truth: dict | None = None,
    session_id: str | None = None,
    run_id: str | None = None,
) -> dict[str, Any]:
    """Shared execution scaffold: build graph, open the per-doc trace (one trace
    per document, deterministic id from `seed`), invoke, emit self-evident
    scores, persist them. Used by both `run_pipeline` and `resume_from_review`.

    Enforces the hard run cutoff: a wall-clock deadline and a cumulative
    output-token budget (see pipeline/limits.py). Aborted runs are finalized to
    the failed bin and still scored + persisted, so every run produces core
    metrics (duration, tokens, cost, call count) for cross-run evaluation.
    """
    import os
    from pipeline import limits
    from observability import tracing

    from observability import scores as pipeline_scores

    graph = build_graph()
    # Attempt-scoped thread: a re-run of the same document must not resume the
    # previous run's checkpointed state (pilot: correspondence_01's degraded
    # second run inherited stale state, producing output=null).
    config = {"configurable": {"thread_id": f"{seed}-run{attempt}"}}

    pipeline_scores.ensure_score_configs()

    started_at = time.time()
    deadline = started_at + float(limits.get_deadline_seconds())
    limits.reset_run_usage()
    limits.set_run_deadline(deadline)
    initial_state = {
        **initial_state,
        "run_deadline": deadline,
        "run_attempt": attempt,
    }
    if ground_truth:
        # Expected outcome for this document (pilot runs pass the manifest
        # ground truth). Carried into the `pipeline-result` generation so the
        # live evaluator can render a CORRECT/PARTIAL/MISS verdict.
        initial_state["ground_truth"] = ground_truth

    # Environment resolution: per-context override (OBSERVABILITY_ENVIRONMENT,
    # set by entrypoints via pipeline.env.default_environment) wins; the
    # standard LANGFUSE_TRACING_ENVIRONMENT is the fallback; mock runs (no
    # observability) are labeled "mock"; everything else defaults to "live".
    environment = (
        os.environ.get("OBSERVABILITY_ENVIRONMENT")
        or os.environ.get("LANGFUSE_TRACING_ENVIRONMENT")
    )
    if not environment:
        if os.environ.get("OBSERVABILITY_PROVIDER", "auto") == "none":
            environment = "mock"
        else:
            environment = "live"

    # Mandatory tag taxonomy (see AGENTS.md "Mandatory: classify and tag every
    # logged run"): `mailroom` (always) + run-context tag matching the
    # environment (`pilot`/`live`/`misc`/`mock`) + attempt tag for re-runs +
    # source corpus tag for pilot/corpus runs.
    tags = ["mailroom", environment]
    if attempt:
        tags.append(f"run-{attempt}")
    if source:
        tags.append(f"source-{source}")

    trace_metadata = {"pipeline": "mailroom", "run_deadline": deadline, "attempt": attempt}
    if source:
        trace_metadata["source"] = source
    if run_id:
        trace_metadata["run_id"] = run_id

    with tracing.pipeline_trace(
        seed=seed,  # deterministic trace id -> correlates with our doc
        session_id=session_id or initial_state.get("matter_id") or "DEFAULT",  # groups documents of a matter/run
        name="document-pipeline",
        input=trace_input,
        metadata=trace_metadata,
        tags=tags,
        environment=environment,
    ) as root:
        # Capture the trace id into the state so manifests, catalog records and
        # the returned result all carry it (the DB↔Langfuse correlation link).
        # get_trace_id() is only valid inside the trace block.
        state_trace_id = tracing.get_trace_id() or ""
        initial_state = {**initial_state, "trace_id": state_trace_id}
        try:
            result = graph.invoke(initial_state, config)
        except (limits.RunDeadlineExceeded, limits.RunBudgetExceeded) as exc:
            logger.warning(
                "run_aborted",
                doc_id=initial_state.get("doc_id"),
                reason=type(exc).__name__,
                detail=str(exc),
            )
            result = _finalize_aborted(initial_state, f"{type(exc).__name__}: {exc}")
        except Exception:
            logger.exception("run_crashed", doc_id=initial_state.get("doc_id"))
            result = _finalize_aborted(initial_state, "unexpected error")
        # Ensure the trace id survives into the final state (ingest_node creates
        # the manifest with its own doc_id; the trace id must be attached even
        # when the graph never ran ingest, e.g. aborted runs).
        if not result.get("trace_id"):
            result["trace_id"] = state_trace_id

        # Post-invoke scoring/emission is best-effort: the file may already be
        # archived/reviewed, so a failure here must never surface as a pipeline
        # failure to the watcher (the run itself succeeded).
        try:
            metrics = pipeline_scores.compute_run_metrics(result, started_at, time.time())
            score_values = pipeline_scores.emit_pipeline_scores(result, metrics)
            _persist_scores(result, score_values)
        except Exception:
            logger.exception("post_run_scoring_failed", doc_id=result.get("doc_id"))
        # Deterministic field-type-aware scoring (issues #4/#5). Only grounded
        # runs have expected field values to compare against; the scorer gates
        # the LLM judge below: an unambiguous deterministic verdict skips the
        # `pipeline-result` generation entirely (2 evaluator calls saved).
        judge_required = None
        expected_fields = (initial_state.get("ground_truth") or {}).get("expected_fields")
        if expected_fields:
            from observability.field_scoring import get_field_types, score_extraction
            from observability.langfuse_field_scoring import score_and_log_extraction

            extracted = result.get("extracted_data") or {}
            doc_class = result.get("doc_type")
            expected_class = (initial_state.get("ground_truth") or {}).get("expected_doc_class")
            if doc_class and extracted:
                try:
                    field_result = score_and_log_extraction(
                        trace_id=tracing.get_trace_id(),
                        doc_class=doc_class,
                        field_types=get_field_types(doc_class),
                        predicted=extracted,
                        expected=expected_fields,
                        matter_id=initial_state.get("matter_id"),
                    )
                    judge_required = field_result.needs_judge_review
                    # A wrong classification must ALWAYS reach the LLM judge:
                    # the deterministic field scorer compares the extraction
                    # against the EXPECTED class's fields, so a misfiled doc
                    # scores ~0 (below the ambiguous band) and would suppress
                    # the verdict for exactly the runs that need scrutiny.
                    if expected_class and doc_class != expected_class:
                        judge_required = True
                        logger.info(
                            "field_scoring_class_mismatch_forces_judge",
                            doc_id=result.get("doc_id"),
                            doc_class=doc_class,
                            expected_class=expected_class,
                        )
                    logger.info(
                        "field_scoring_computed",
                        doc_id=result.get("doc_id"),
                        overall=field_result.overall_score,
                        ambiguous_fields=field_result.ambiguous_fields,
                        judge_required=judge_required,
                    )
                except Exception:
                    logger.exception("field_scoring_failed")
        try:
            if root is not None:
                _emit_pipeline_result(root, result, initial_state, judge_required=judge_required)
                root.update(output={
                    "stage": result.get("stage"),
                    "doc_type": result.get("doc_type"),
                    "classification_confidence": result.get("classification_confidence"),
                    "extraction_confidence": result.get("extraction_confidence"),
                    "run_aborted": bool(result.get("run_aborted")),
                    "error_message": result.get("error_message"),
                })
        except Exception:
            logger.exception("pipeline_result_emission_failed", doc_id=result.get("doc_id"))
    tracing.flush()
    return result


def run_pipeline(
    file_path: Path,
    matter_id: str = "DEFAULT",
    attempt: int = 0,
    source: str | None = None,
    ground_truth: dict | None = None,
    session_id: str | None = None,
    run_id: str | None = None,
) -> dict[str, Any]:
    _ensure_dirs()

    initial_state: DocumentState = {
        "doc_id": "",
        "matter_id": matter_id,
        "original_filename": file_path.name,
        "stage": "inbox",
        "doc_type": None,
        "classification_confidence": None,
        "classification_attempts": 0,
        "extracted_data": None,
        "extraction_confidence": None,
        "extraction_attempts": 0,
        "trace_id": None,
        "escalation_reason": None,
        "review_decision": None,
        "retry_count": 0,
        "conflict_detected": False,
        "file_path": str(file_path),
        "doc_text": "",
        "doc_pages": [],
        "error_message": None,
        "messages": [],
        "transient_error": False,
        "transient_retries": 0,
        "run_attempt": attempt,
    }

    # Attempt 0 keeps the bare filename stem as the deterministic trace seed
    # (backwards-compatible with ground-truth score ingestion in run_pilot.py);
    # subsequent attempts (e.g. scheduled re-processing) get a suffixed seed so
    # each run gets its own trace instead of merging into one misleading span.
    seed = file_path.stem if attempt <= 0 else f"{file_path.stem}-run{attempt}"

    return _execute_run(
        initial_state,
        seed=seed,
        attempt=attempt,
        source=source,
        ground_truth=ground_truth,
        session_id=session_id,
        run_id=run_id,
        trace_input={"filename": file_path.name, "matter_id": matter_id, "attempt": attempt},
    )


def resume_from_review(manifest, review_file: Path) -> dict[str, Any]:
    """Resume a human-approved review document with a FRESH extraction.

    Stateless resume: the file is requeued from the review bin into
    processing/<worker>/, re-read, and the graph is re-invoked starting at the
    extraction stage (skipping classification — the manifest's doc_type is
    trusted from the reviewed run). The ORIGINAL doc_id is preserved so the
    manifest, catalog record, audit chain, and Langfuse trace stay intact.
    """
    from pipeline.bins import requeue_from_review, get_worker_id

    if not manifest.doc_type:
        raise ValueError(
            "Cannot resume: manifest has no classification; re-submit the document instead."
        )

    _ensure_dirs()
    worker_id = get_worker_id()
    queued = requeue_from_review(review_file, worker_id)
    doc_text, text_ok = _read_file_text(queued)
    doc_pages = _render_doc_pages(queued)

    initial_state: DocumentState = {
        "doc_id": manifest.doc_id,
        "matter_id": manifest.matter_id,
        "original_filename": manifest.original_filename,
        "stage": PipelineStage.CLASSIFIED.value,
        "doc_type": manifest.doc_type,
        "contract_subtype": manifest.contract_subtype,
        "classification_confidence": manifest.classification_confidence,
        "classification_attempts": manifest.classification_attempts,
        "extracted_data": None,  # fresh extraction — never reuse the reviewed data
        "extraction_confidence": None,
        "extraction_attempts": 0,
        "trace_id": None,
        "escalation_reason": None,
        "review_decision": "approved",
        "retry_count": 0,
        "conflict_detected": False,
        "file_path": str(queued),
        "doc_text": doc_text,
        "doc_pages": doc_pages,
        "error_message": None if text_ok else f"Could not extract text from {queued.suffix} file",
        "messages": [],
        "resume_extraction": True,
        "transient_error": False,
        "transient_retries": 0,
        "run_attempt": 0,
    }

    return _execute_run(
        initial_state,
        seed=queued.stem,
        attempt=0,
        trace_input={"filename": queued.name, "matter_id": manifest.matter_id, "resumed": True},
    )
